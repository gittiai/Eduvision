import streamlit as st
from PIL import Image
from pdf2image import convert_from_bytes
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from docx import Document
import os
from dotenv import load_dotenv
import google.generativeai as genai

st.markdown(
    """
    <style>
    section[data-testid="stSidebar"] {
        background-image: url("https://images.unsplash.com/photo-1611325058416-db7794e8e32c?q=80&w=2070&auto=format&fit=crop&ixlib=rb-4.1.0&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D");

        background-size: cover;
        background-repeat: no-repeat;
    }
    </style>
    """,
    unsafe_allow_html=True
)
st.markdown(
    """
    <style>
    .stApp {
        background-image: url("https://images.unsplash.com/photo-1711062717319-393e424a3538?q=80&w=2880&auto=format&fit=crop&ixlib=rb-4.1.0&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ----------------- CONFIG -----------------
load_dotenv()
gemini_api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=gemini_api_key)
os.environ["GOOGLE_API_KEY"] = os.getenv("GEMINI_API_KEY")
# ----------------- Helper: Extract text from file -----------------
def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pages = convert_from_bytes(uploaded_file.read(), dpi=300)
        images = [page for page in pages]
    else:
        images = [Image.open(uploaded_file)]

    model = genai.GenerativeModel("gemini-1.5-flash")
    extracted_text = ""
    for img in images:
        response = model.generate_content(
            ["Extract all useful and meaningful text (especially math/equations) from this image.", img],
            stream=False
        )
        extracted_text += response.text + "\n\n"
    return extracted_text, images

# ----------------- Streamlit Tabs -----------------
st.title("📝 EduVision:Your Learning Partner")

tab1, tab2 = st.tabs(["📖 OCR + Cleaning + QA", "📚 Paper Evaluation"])

# ----------------- TAB 1 : OCR + Cleaning + QA -----------------
with tab1:
    uploaded_file = st.file_uploader("📄 Upload handwritten PDF/Image", type=["jpg","jpeg","png","pdf"])
    question = st.text_input("💬 Ask a question about this document (optional)")

    if uploaded_file:
        # Extract text
        raw_text, images = extract_text_from_file(uploaded_file)

        # Show preview images
        for img in images:
            st.image(img, use_container_width=True)

        # Clean extracted text
        st.info("Cleaning extracted text with AI...")
        llm_cleaner = ChatGroq(model="openai/gpt-oss-20b")
        prompt_clean = f"""
You are an expert at converting handwritten math text into a clean, readable, and well-formatted digital version. 
Preserve **all mathematical expressions, symbols, equations, fractions, indices, superscripts, subscripts, and operators** exactly as they appear. 
Do not change numbers, variable names, or operators. Keep fractions as 'a/b', exponents as 'x^2', and maintain line structure for clarity.

Here is the extracted text from OCR/GenAI:
{raw_text}

Return a cleaned, readable version with all math symbols preserved exactly.
"""
        cleaned_text = llm_cleaner.predict(prompt_clean)

        st.subheader("✨ Cleaned Text")
        st.text_area("Cleaned output", cleaned_text, height=300)

        # ----------------- Optional QA -----------------
        if question:
            st.info("Running QA on cleaned text...")
            splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
            docs = splitter.split_text(cleaned_text)

            #embeddings = OllamaEmbeddings(model="nomic-embed-text")

            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
            vectordb = FAISS.from_texts(docs, embedding=embeddings)

            llm_qa = ChatGroq(model="openai/gpt-oss-20b")
            retriever = vectordb.as_retriever()
            qa_chain = RetrievalQA.from_chain_type(llm=llm_qa, retriever=retriever)

            raw_answer = qa_chain.run(question)

            # Format answer with math preservation
            st.info("Formatting QA answer with math preservation...")
            prompt_template = ChatPromptTemplate.from_template("""
You are an expert at presenting math answers clearly. 
Rules:
- Use **real math symbols** (∫, √, ÷, ×, ≤, ≥, etc.).
- For exponents use Unicode (x², x³, …), not ^.
- For fractions, use "a/b" format instead of LaTeX \frac.
- Never wrap output in LaTeX delimiters like \\[ ... \\].
- Keep line breaks for steps (Step 1, Step 2, etc.).
- Output must be human-readable in plain text.

Question: {question}
Answer (raw, may contain LaTeX): {raw_answer}

Now return a clean, step-by-step solution with correct math symbols (no LaTeX).
""")


            final_prompt = prompt_template.format(question=question, raw_answer=raw_answer)
            formatted_answer = llm_qa.predict(final_prompt)

            st.subheader("💡 AI Answer (Math Preserved)")
            st.text_area("Final Answer", formatted_answer, height=250)

        # Save as Word
        doc = Document()
        doc.add_paragraph(cleaned_text)
        doc.save("output.docx")
        st.download_button("⬇️ Download Word File", open("output.docx", "rb"), "output.docx")

# ----------------- TAB 2 : Paper Evaluation -----------------
with tab2:
    q_file = st.file_uploader("📄 Upload Question Paper", type=["jpg","jpeg","png","pdf"], key="qfile")
    a_file = st.file_uploader("📝 Upload Student Answer Sheet", type=["jpg","jpeg","png","pdf"], key="afile")

    if q_file and a_file:
        st.info("Processing question paper...")
        question_text, _ = extract_text_from_file(q_file)
        st.text_area("📜 Extracted Questions", question_text, height=200)

        st.info("Processing answer sheet...")
        answer_text, _ = extract_text_from_file(a_file)
        st.text_area("📝 Extracted Answers", answer_text, height=200)

        # Evaluation
        st.info("Evaluating answers with AI...")
        llm_evaluator = ChatGroq(model="openai/gpt-oss-20b")

        prompt_eval = f"""
You are an exam evaluator. 
Here are the exam questions:

{question_text}

Here are the student's answers:

{answer_text}

Please evaluate the student's answers:
- Check correctness of each answer.
- Give feedback on mistakes and improvements.
- Suggest marks/score if possible (assume each question is equal marks).
- Keep math symbols as written (do not convert to LaTeX).
Return a clear evaluation report.
"""
        evaluation = llm_evaluator.predict(prompt_eval)

        st.subheader("📊 Evaluation Report")
        st.write(evaluation)

        # Save as Word
        doc = Document()
        doc.add_heading("AI Evaluation Report", level=1)
        doc.add_paragraph(evaluation)
        doc.save("evaluation_report.docx")
        st.download_button("⬇️ Download Evaluation Report", open("evaluation_report.docx", "rb"), "evaluation_report.docx")
